"""
Document Processing Service - Extract text from documents and generate prompts.

Handles:
- PDF text extraction using pypdf
- DOCX text extraction using python-docx
- LLM-based prompt generation from extracted text
"""
import hashlib
import io
import logging
from typing import Optional, List, Dict, Any

from django.conf import settings
from django.utils import timezone

logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Error during document processing."""
    pass


class DocumentProcessor:
    """
    Process uploaded policy documents and extract prompts.
    """

    # Prompt extraction system message
    EXTRACTION_SYSTEM_PROMPT = """You are an expert at analyzing company policy documents and extracting actionable guidelines for AI systems.

Your task is to analyze the provided policy document and extract specific rules, constraints, and guidelines that should be enforced by an AI assistant.

For each identified policy area, generate a structured prompt that can be used as part of an AI system's instructions.

Categories to look for:
- SAFETY: Guidelines about harmful content, sensitive topics, safety measures
- CONSTRAINTS: Rules about what the AI can/cannot do, limitations
- PERSONA: Expected tone, personality, communication style
- CONTEXT: Background information, domain knowledge to maintain
- COMPLIANCE: Regulatory requirements (HIPAA, GDPR, SOC2, etc.)

Output Format (JSON):
{
  "summary": "Brief summary of the document's purpose",
  "detected_categories": ["SAFETY", "CONSTRAINTS", ...],
  "prompts": [
    {
      "category": "SAFETY",
      "title": "Short descriptive title",
      "prompt_text": "The actual prompt instruction...",
      "source_excerpt": "Relevant quote from document",
      "confidence": 0.95
    }
  ],
  "compliance_frameworks": ["HIPAA", "SOC2", ...],
  "risk_level": "low|medium|high"
}"""

    def __init__(self, organization=None):
        self.organization = organization

    def extract_text_from_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text from a PDF file.

        Args:
            file_content: Raw PDF file bytes

        Returns:
            Dict with 'text', 'page_count', 'word_count'
        """
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(file_content))
            pages = []

            for page in reader.pages:
                text = page.extract_text() or ''
                pages.append(text)

            full_text = '\n\n'.join(pages)
            word_count = len(full_text.split())

            return {
                'text': full_text,
                'page_count': len(reader.pages),
                'word_count': word_count,
            }
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            raise DocumentProcessingError(f"PDF extraction failed: {e}")

    def extract_text_from_docx(self, file_content: bytes) -> Dict[str, Any]:
        """
        Extract text from a DOCX file.

        Args:
            file_content: Raw DOCX file bytes

        Returns:
            Dict with 'text', 'page_count', 'word_count'
        """
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_content))
            paragraphs = []

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            full_text = '\n\n'.join(paragraphs)
            word_count = len(full_text.split())

            # DOCX doesn't have page concept like PDF
            # Estimate ~500 words per page
            estimated_pages = max(1, word_count // 500)

            return {
                'text': full_text,
                'page_count': estimated_pages,
                'word_count': word_count,
            }
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            raise DocumentProcessingError(f"DOCX extraction failed: {e}")

    def extract_text_from_txt(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from a plain text file."""
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                text = file_content.decode('latin-1')

            word_count = len(text.split())
            estimated_pages = max(1, word_count // 500)

            return {
                'text': text,
                'page_count': estimated_pages,
                'word_count': word_count,
            }
        except Exception as e:
            logger.error(f"Failed to extract TXT text: {e}")
            raise DocumentProcessingError(f"TXT extraction failed: {e}")

    def extract_text(self, file_content: bytes, document_type: str) -> Dict[str, Any]:
        """
        Extract text from a document based on its type.

        Args:
            file_content: Raw file bytes
            document_type: 'pdf', 'docx', or 'txt'

        Returns:
            Dict with 'text', 'page_count', 'word_count'
        """
        extractors = {
            'pdf': self.extract_text_from_pdf,
            'docx': self.extract_text_from_docx,
            'txt': self.extract_text_from_txt,
        }

        extractor = extractors.get(document_type.lower())
        if not extractor:
            raise DocumentProcessingError(f"Unsupported document type: {document_type}")

        return extractor(file_content)

    def compute_file_hash(self, file_content: bytes) -> str:
        """Compute SHA-256 hash of file content."""
        return hashlib.sha256(file_content).hexdigest()

    async def analyze_with_llm(
        self,
        text: str,
        model: str = 'gpt-4o-mini',
    ) -> Dict[str, Any]:
        """
        Use LLM to analyze extracted text and generate prompts.

        Args:
            text: Extracted document text
            model: LLM model to use

        Returns:
            Analysis results with generated prompts
        """
        import json
        from zentinelle.services.ai_service import AIService

        # Truncate text if too long (keep first ~100k chars)
        max_chars = 100000
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n[Document truncated for analysis...]"

        ai_service = AIService(organization=self.organization)

        try:
            response = await ai_service.chat_completion(
                model=model,
                messages=[
                    {'role': 'system', 'content': self.EXTRACTION_SYSTEM_PROMPT},
                    {'role': 'user', 'content': f"Analyze this policy document and extract prompts:\n\n{text}"},
                ],
                response_format={'type': 'json_object'},
                temperature=0.3,
            )

            # Parse JSON response
            content = response.get('choices', [{}])[0].get('message', {}).get('content', '{}')
            return json.loads(content)

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM response: {e}")
            return {
                'summary': 'Failed to parse analysis',
                'detected_categories': [],
                'prompts': [],
                'error': str(e),
            }
        except Exception as e:
            logger.error(f"LLM analysis failed: {e}")
            return {
                'summary': 'Analysis failed',
                'detected_categories': [],
                'prompts': [],
                'error': str(e),
            }

    def create_prompts_from_analysis(
        self,
        policy_document,
        analysis: Dict[str, Any],
    ) -> List:
        """
        Create SystemPrompt objects from LLM analysis.

        Args:
            policy_document: PolicyDocument instance
            analysis: LLM analysis results

        Returns:
            List of created SystemPrompt objects
        """
        from zentinelle.models import SystemPrompt
        from django.utils.text import slugify

        created_prompts = []
        prompts_data = analysis.get('prompts', [])

        # Map categories to SystemPrompt.PromptType
        category_map = {
            'SAFETY': SystemPrompt.PromptType.SAFETY,
            'CONSTRAINTS': SystemPrompt.PromptType.CONSTRAINTS,
            'PERSONA': SystemPrompt.PromptType.PERSONA,
            'CONTEXT': SystemPrompt.PromptType.CONTEXT,
            'COMPLIANCE': SystemPrompt.PromptType.CONSTRAINTS,
            'BASE': SystemPrompt.PromptType.BASE,
        }

        for i, prompt_data in enumerate(prompts_data):
            category = prompt_data.get('category', 'CUSTOM').upper()
            prompt_type = category_map.get(category, SystemPrompt.PromptType.CUSTOM)

            title = prompt_data.get('title', f'Extracted Prompt {i + 1}')
            slug_base = slugify(title)[:80]

            # Create the prompt
            prompt = SystemPrompt.objects.create(
                organization=policy_document.organization,
                name=title,
                slug=f"{slug_base}-{str(policy_document.id)[:8]}",
                description=f"Extracted from: {policy_document.name}\n\nSource: {prompt_data.get('source_excerpt', 'N/A')[:500]}",
                prompt_text=prompt_data.get('prompt_text', ''),
                prompt_type=prompt_type,
                status=SystemPrompt.Status.DRAFT,
                is_template=False,
                metadata={
                    'source_document_id': str(policy_document.id),
                    'confidence': prompt_data.get('confidence', 0),
                    'extracted_at': timezone.now().isoformat(),
                },
            )
            created_prompts.append(prompt)

        return created_prompts


class PolicyDocumentService:
    """
    High-level service for managing policy document uploads and processing.
    """

    def __init__(self, organization=None):
        self.organization = organization
        self.processor = DocumentProcessor(organization=organization)

    def upload_document(
        self,
        name: str,
        file_content: bytes,
        document_type: str,
        description: str = '',
    ):
        """
        Upload and queue a policy document for processing.

        Args:
            name: Document name
            file_content: Raw file bytes
            document_type: 'pdf', 'docx', or 'txt'
            description: Optional description

        Returns:
            Created PolicyDocument instance
        """
        from zentinelle.models import PolicyDocument

        # Compute hash for deduplication
        file_hash = self.processor.compute_file_hash(file_content)

        # Check for existing document with same hash
        existing = PolicyDocument.objects.filter(
            organization=self.organization,
            file_hash=file_hash,
        ).first()

        if existing:
            logger.info(f"Document with hash {file_hash[:16]} already exists: {existing.id}")
            return existing

        # Store file (in-memory for now, would upload to S3 in production)
        # For now, we'll process synchronously
        file_path = f"policy-docs/{self.organization.id}/{file_hash[:16]}.{document_type}"

        document = PolicyDocument.objects.create(
            organization=self.organization,
            name=name,
            description=description,
            document_type=document_type,
            file_path=file_path,
            file_size=len(file_content),
            file_hash=file_hash,
            status=PolicyDocument.Status.PENDING,
        )

        # Extract text immediately
        try:
            extraction = self.processor.extract_text(file_content, document_type)
            document.extracted_text = extraction['text']
            document.page_count = extraction['page_count']
            document.word_count = extraction['word_count']
            document.status = PolicyDocument.Status.EXTRACTED
            document.save()
        except DocumentProcessingError as e:
            document.status = PolicyDocument.Status.FAILED
            document.error_message = str(e)
            document.save()

        return document

    async def analyze_document(
        self,
        document,
        model: str = 'gpt-4o-mini',
    ) -> Dict[str, Any]:
        """
        Run LLM analysis on an extracted document.

        Args:
            document: PolicyDocument instance
            model: LLM model to use

        Returns:
            Analysis results
        """
        if document.status != document.Status.EXTRACTED:
            raise DocumentProcessingError(
                f"Document must be in EXTRACTED status, got {document.status}"
            )

        document.status = document.Status.PROCESSING
        document.save(update_fields=['status'])

        try:
            # Run LLM analysis
            analysis = await self.processor.analyze_with_llm(
                document.extracted_text,
                model=model,
            )

            # Store results
            document.analysis_results = analysis
            document.processing_model = model
            document.processed_at = timezone.now()

            # Create prompts
            prompts = self.processor.create_prompts_from_analysis(document, analysis)

            # Link prompts to document
            document.generated_prompts.set(prompts)
            document.status = document.Status.ANALYZED
            document.save()

            return analysis

        except Exception as e:
            document.status = document.Status.FAILED
            document.error_message = str(e)
            document.save()
            raise
